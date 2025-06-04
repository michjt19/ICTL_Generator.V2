from docx import Document

def generate_word(task, filepath):
    doc = Document()
    doc.add_heading(task["title"], level=1)
    doc.add_paragraph(f"Task Code: {task['code']}")
    doc.add_paragraph(f"Condition: {task['condition']}")
    doc.add_paragraph(f"Standard: {task['standard']}")

    if "steps" in task:
        doc.add_heading("Performance Steps", level=2)
        for step in task["steps"]:
            doc.add_paragraph(step, style="List Number")

    doc.save(filepath)
